import json
from docx import Document

def generate_report(listener, analyst, vc, critic, advisor, out_file="Idea_Report.docx"):
    doc = Document()
    doc.add_heading("Idea Validation Report", 0)

    # Section 1: Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(f"Title: {listener.get('idea_title')}")
    doc.add_paragraph(f"One-liner: {listener.get('one_liner')}")
    doc.add_paragraph(f"Problem: {listener.get('problem')}")
    doc.add_paragraph(f"Solution: {listener.get('solution')}")
    doc.add_paragraph(f"Target Customer: {listener.get('target_customer')}")

    # Section 2: Market Analysis
    doc.add_heading("2. Market Analysis", level=1)
    doc.add_paragraph(json.dumps(analyst, indent=2, ensure_ascii=False))

    # Section 3: VC Feasibility
    doc.add_heading("3. Feasibility Assessment", level=1)
    doc.add_paragraph(json.dumps(vc, indent=2, ensure_ascii=False))

    # Section 4: Challenges
    doc.add_heading("4. Challenges & Blind Spots", level=1)
    doc.add_paragraph(json.dumps(critic, indent=2, ensure_ascii=False))

    # Section 5: Advisor Plan
    doc.add_heading("5. Action Plan", level=1)
    doc.add_paragraph(json.dumps(advisor, indent=2, ensure_ascii=False))

    doc.save(out_file)
    print(f"✅ Report generated: {out_file}")
